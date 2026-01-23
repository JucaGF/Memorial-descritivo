"""Montagem do documento DOCX final."""

from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

from docx import Document

from memorial_maker.writer.docx_styles import (
    setup_document_margins,
    setup_styles,
    add_header_footer,
    add_section_heading,
    add_body_text,
    add_signature_block,
)
from memorial_maker.utils.logging import get_logger

logger = get_logger("writer.docx")


class MemorialWriter:
    """Escritor de Memorial em DOCX."""

    def __init__(self):
        """Inicializa writer.
        
        """
        self.doc = Document()
        
        # Configura margens conforme modelo
        setup_document_margins(self.doc)
        
        # Configura estilos
        setup_styles(self.doc)

    def write_memorial(
        self,
        sections: Dict[str, str],
        master_data: Dict,
        output_path: Path,
        memorial_type: str = "telecom",
    ):
        """Escreve memorial completo.
        
        Args:
            sections: Dicionário {section_id: content}
            master_data: Dados consolidados (JSON mestre)
            output_path: Caminho de saída do DOCX
            memorial_type: Tipo de memorial ("telecom" ou "eletrico")
        """
        logger.info(f"Escrevendo memorial: {output_path} (tipo: {memorial_type})")
        
        # Debug: verifica estrutura do master_data
        logger.info(f"Master data keys: {list(master_data.keys())}")
        if "obra" in master_data:
            logger.info(f"Obra keys: {list(master_data['obra'].keys())}")
        
        # Dados do projeto
        project_data = {
            "empreendimento": master_data.get("obra", {}).get("empreendimento", ""),
            "construtora": master_data.get("obra", {}).get("construtora", ""),
            "endereco": master_data.get("obra", {}).get("endereco", ""),
            "carimbo": master_data.get("obra", {}).get("carimbo", {}),
        }
        
        logger.info(f"Project data: {project_data}")
        
        # Adiciona cabeçalho e rodapé (para todas as páginas)
        add_header_footer(self.doc, project_data)
        
        # Capa do memorial
        from memorial_maker.writer.docx_styles import add_cover_page
        add_cover_page(self.doc, None, project_data, memorial_type)
        
        # Sumário - conditional based on memorial type
        if memorial_type == "eletrico":
            # For electrical memorials, the TOC is generated as s1_sumario section
            # and will be written by _write_electrical_sections()
            pass
        else:
            # For telecom memorials, use the hardcoded TOC
            from memorial_maker.writer.docx_styles import add_table_of_contents
            add_table_of_contents(self.doc)
        
        
        # Seções baseadas no tipo
        if memorial_type == "eletrico":
            self._write_electrical_sections(sections)
        else:
            # Telecom memorial (existing behavior)
            self._write_section_1(sections.get("s1_introducao", ""))
            self._write_section_2(sections.get("s2_dados_obra", ""))
            self._write_section_3(sections.get("s3_normas", ""))
            self._write_section_4(sections)
            self._write_section_5(sections.get("s5_sala_monitoramento", ""))
            self._write_section_6(sections.get("s6_passivos_ativos", ""))
            self._write_section_7(sections.get("s7_testes_aceitacao", ""))
            
            # Adiciona bloco de assinatura no final do memorial telecom
            add_signature_block(self.doc, master_data)
        
        # Salva
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        logger.info(f"Memorial salvo: {output_path}")

    def _write_section_1(self, content: str):
        """Seção 1: Introdução."""
        add_section_heading(self.doc, "1", "INTRODUÇÃO", level=1)
        add_body_text(self.doc, content)

    def _write_section_2(self, content: str):
        """Seção 2: Dados da Obra."""
        add_section_heading(self.doc, "2", "DADOS DA OBRA", level=1)
        # Use LEFT alignment for better readability with bullet lists
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            para = self.doc.add_paragraph(para_text, style='Normal')
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _write_section_3(self, content: str):
        """Seção 3: Normas Técnicas."""
        add_section_heading(self.doc, "3", "NORMAS TÉCNICAS", level=1)
        add_body_text(self.doc, content)

    def _write_section_4(self, sections: Dict[str, str]):
        """Seção 4: Serviços Contemplados (com subseções)."""
        add_section_heading(self.doc, "4", "SERVIÇOS CONTEMPLADOS", level=1)
        
        # Introdução da seção 4
        intro = sections.get("s4_servicos", "")
        add_body_text(self.doc, intro)
        
        # Subseções
        subsections = [
            ("4.1", "SERVIÇO DE VOZ", "s4_1_voz"),
            ("4.2", "SERVIÇO DE DADOS", "s4_2_dados"),
            ("4.3", "SERVIÇO DE VÍDEO", "s4_3_video"),
            ("4.4", "SERVIÇO DE INTERCOMUNICAÇÃO", "s4_4_intercom"),
            ("4.5", "SERVIÇO DE MONITORAMENTO", "s4_5_monitoramento"),
        ]
        
        for number, title, section_id in subsections:
            content = sections.get(section_id, "")
            if content:
                add_section_heading(self.doc, number, title, level=2)
                add_body_text(self.doc, content)

    def _write_section_5(self, content: str):
        """Seção 5: Sala de Monitoramento."""
        add_section_heading(self.doc, "5", "SALA DE MONITORAMENTO (ER/EF)", level=1)
        add_body_text(self.doc, content)

    def _write_section_6(self, content: str):
        """Seção 6: Elementos Passivos e Ativos."""
        add_section_heading(self.doc, "6", "ELEMENTOS PASSIVOS E ATIVOS DA REDE", level=1)
        add_body_text(self.doc, content)

    def _write_section_7(self, content: str):
        """Seção 7: Testes e Aceitação."""
        add_section_heading(self.doc, "7", "TESTES E ACEITAÇÃO", level=1)
        add_body_text(self.doc, content)

    def _write_electrical_sections(self, sections: Dict[str, str]):
        """Escreve seções do memorial elétrico (dinâmico, apenas seções presentes).
        
        Matches reference memorial structure with sections 1-5.
        
        Args:
            sections: Dicionário {section_id: content} - apenas seções com conteúdo
        """
        # Sumário (Static TOC)
        if "s0_sumario" in sections:
            content = sections["s0_sumario"]
            if content and len(content.strip()) > 50:
                add_section_heading(self.doc, "", "SUMÁRIO", level=1)
                # Add sumário with LEFT alignment to avoid excessive word spacing
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                paragraphs = content.split("\n\n")
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if not para_text:
                        continue
                    para = self.doc.add_paragraph(para_text, style='Normal')
                    # Override alignment to LEFT for better readability in TOC
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Section 1: INTRODUÇÃO
        if "s1_introducao" in sections:
            content = sections["s1_introducao"]
            if content and len(content.strip()) > 50:
                add_section_heading(self.doc, "1", "INTRODUÇÃO", level=1)
                add_body_text(self.doc, content)
        
        # Section 2: DADOS DA OBRA
        if "s2_dados_obra" in sections:
            content = sections["s2_dados_obra"]
            if content and len(content.strip()) > 50:
                add_section_heading(self.doc, "2", "DADOS DA OBRA", level=1)
                # Use LEFT alignment for better readability with bullet lists
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                paragraphs = content.split("\n\n")
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if not para_text:
                        continue
                    para = self.doc.add_paragraph(para_text, style='Normal')
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Section 3: REQUISITOS GERAIS
        if "s3_requisitos_gerais" in sections:
            content = sections["s3_requisitos_gerais"]
            if content and len(content.strip()) > 50:
                add_section_heading(self.doc, "3", "REQUISITOS GERAIS", level=1)
                add_body_text(self.doc, content)
        
        # Section 3.1: DISPOSIÇÕES GERAIS
        if "s3_1_disposicoes_gerais" in sections:
            content = sections["s3_1_disposicoes_gerais"]
            if content and len(content.strip()) > 50:
                add_section_heading(self.doc, "3.1", "DISPOSIÇÕES GERAIS", level=2)
                add_body_text(self.doc, content)
        
        # Section 3.2: DISPOSIÇÕES DE BASE DE PROJETO
        if "s3_2_disposicoes_base_projeto" in sections:
            content = sections["s3_2_disposicoes_base_projeto"]
            if content and len(content.strip()) > 50:
                add_section_heading(self.doc, "3.2", "DISPOSIÇÕES DE BASE DE PROJETO", level=2)
                add_body_text(self.doc, content)
        
        # Section 4: VISÃO GERAL DO SISTEMA DE INSTALAÇÕES ELÉTRICAS
        if "s4_visao_geral" in sections:
            content = sections["s4_visao_geral"]
            if content and len(content.strip()) > 50:
                add_section_heading(self.doc, "4", "VISÃO GERAL DO SISTEMA DE INSTALAÇÕES ELÉTRICAS", level=1)
                add_body_text(self.doc, content)
        
        # Section 4 subsections (4.1-4.12) - conditional
        section4_subsections = [
            ("4.1", "INSTALAÇÃO DE ENTRADA E MEDIÇÃO DE ENERGIA", "s4_1_entrada_energia"),
            ("4.2", "SUBESTAÇÃO ABRIGADA", "s4_2_subestacao"),
            ("4.2.1", "DADOS CONSTRUTIVOS DA SUBESTAÇÃO", "s4_2_1_dados_construtivos"),
            ("4.3", "ATERRAMENTO", "s4_3_aterramento"),
            ("4.4", "ELETRODUTOS, LEITOS E CANALETAS", "s4_4_eletrodutos_leitos"),
            ("4.5", "CONDUTORES", "s4_5_condutores"),
            ("4.6", "MEDIÇÃO DE ENERGIA ELÉTRICA", "s4_6_medicao_energia"),
            ("4.7", "CORES PADRONIZADAS", "s4_7_cores"),
            ("4.8", "INSTALAÇÃO DE LUZ E FORÇA", "s4_8_luz_forca"),
            ("4.9", "INSTALAÇÃO DE LUZ ESSENCIAL (GRUPO GERADOR)", "s4_9_luz_essencial"),
            ("4.10", "PROTEÇÃO CONTRA CHOQUES ELÉTRICOS E ATERRAMENTO", "s4_10_protecao_aterramento"),
            ("4.11", "MONTAGEM DOS APARELHOS", "s4_11_montagem_aparelhos"),
            ("4.12", "ITENS NÃO INCLUSOS", "s4_12_itens_nao_inclusos"),
        ]
        
        for number, title, section_id in section4_subsections:
            if section_id in sections:
                content = sections[section_id]
                if content and len(content.strip()) > 50:
                    # Determine level: 4.2.1 is level 3, others are level 2
                    level = 3 if number == "4.2.1" else 2
                    add_section_heading(self.doc, number, title, level=level)
                    add_body_text(self.doc, content)
        
        # Section 5: ESPECIFICAÇÕES DOS MATERIAIS
        if "s5_especificacao_materiais" in sections:
            content = sections["s5_especificacao_materiais"]
            if content and len(content.strip()) > 50:
                add_section_heading(self.doc, "5", "ESPECIFICAÇÕES DOS MATERIAIS", level=1)
                add_body_text(self.doc, content)
        
        # Section 5 subsections (5.1-5.4) - always included
        section5_subsections = [
            ("5.1", "INSTALAÇÕES ELÉTRICAS", "s5_1_instalacoes_eletricas"),
            ("5.2", "FIOS E CABOS DE ENERGIA E COMANDO", "s5_2_fios_cabos"),
            ("5.3", "LUMINÁRIAS E ACESSÓRIOS", "s5_3_luminarias"),
            ("5.4", "QUADROS GERAIS DE LUZ E FORÇA E ACESSÓRIOS", "s5_4_quadros"),
        ]
        
        for number, title, section_id in section5_subsections:
            if section_id in sections:
                content = sections[section_id]
                if content and len(content.strip()) > 50:
                    add_section_heading(self.doc, number, title, level=2)
                    add_body_text(self.doc, content)


def write_memorial_docx(
    sections: Dict[str, str],
    master_data: Dict,
    output_dir: Path,
    logo_path: Optional[Path] = None,
    project_name: str = "PROJETO",
    memorial_type: str = "telecom",
) -> Path:
    """Função conveniente para escrever memorial.
    
    Args:
        sections: Seções geradas
        master_data: Dados consolidados
        output_dir: Diretório de saída
        logo_path: Caminho para logo
        project_name: Nome do projeto
        memorial_type: Tipo de memorial ("telecom" ou "eletrico")
        
    Returns:
        Caminho do arquivo gerado
    """
    timestamp = datetime.now().strftime("%Y-%m-%d")
    memorial_type_label = "ELETRICO" if memorial_type == "eletrico" else "TELECOM"
    filename = f"MEMORIAL_{memorial_type_label}_{project_name}_{timestamp}.docx"
    output_path = output_dir / filename
    
    writer = MemorialWriter()
    writer.write_memorial(sections, master_data, output_path, memorial_type=memorial_type)
    
    return output_path











