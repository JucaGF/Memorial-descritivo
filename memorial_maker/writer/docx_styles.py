"""Estilos e formatação para DOCX."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def setup_document_margins(doc: Document):
    """Configura margens do documento conforme modelo.
    
    Margens do modelo:
    - Superior: 3.5 cm
    - Inferior: 2.54 cm
    - Esquerda: 2.5 cm
    - Direita: 2.0 cm
    """
    for section in doc.sections:
        section.top_margin = Cm(3.5)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        # Cabeçalho e rodapé bem nas extremidades
        section.header_distance = Cm(0)  # Cola no topo
        section.footer_distance = Cm(0)  # Cola no fundo


def setup_styles(doc: Document):
    """Configura estilos do documento conforme modelo.
    
    Args:
        doc: Documento python-docx
    """
    
    # Estilo: Título 1 (seções principais: 1, 2, 3, ...)
    try:
        style_h1 = doc.styles['Heading 1']
    except KeyError:
        style_h1 = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    
    style_h1.font.name = 'Arial'
    style_h1.font.size = Pt(12)
    style_h1.font.bold = True
    # Remove cor específica para usar a cor padrão do texto
    style_h1.paragraph_format.space_before = Pt(12)
    style_h1.paragraph_format.space_after = Pt(6)
    
    # Estilo: Título 2 (subseções: 4.1, 4.2, ...)
    try:
        style_h2 = doc.styles['Heading 2']
    except KeyError:
        style_h2 = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    
    style_h2.font.name = 'Arial'
    style_h2.font.size = Pt(11)
    style_h2.font.bold = True
    # Remove cor específica para usar a cor padrão do texto
    style_h2.paragraph_format.space_before = Pt(10)
    style_h2.paragraph_format.space_after = Pt(6)
    
    # Estilo: Corpo (texto normal)
    try:
        style_body = doc.styles['Normal']
    except KeyError:
        style_body = doc.styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
    
    style_body.font.name = 'Arial'
    style_body.font.size = Pt(11)
    style_body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_body.paragraph_format.space_after = Pt(6)
    style_body.paragraph_format.line_spacing = 1.15
    
    # Estilo: Lista
    try:
        style_list = doc.styles['List Bullet']
    except KeyError:
        style_list = doc.styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
    
    style_list.font.name = 'Arial'
    style_list.font.size = Pt(11)
    style_list.paragraph_format.left_indent = Inches(0.25)
    style_list.paragraph_format.space_after = Pt(4)


def add_header_footer(doc: Document, project_data: dict):
    """Adiciona cabeçalho e rodapé com imagens esticadas (largura total da página).
    
    Args:
        doc: Documento python-docx
        project_data: Dados do projeto
    """
    from pathlib import Path
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    section = doc.sections[0]
    
    # Caminhos das imagens ORIGINAIS (sem redimensionamento prévio)
    assets_dir = Path(__file__).parent.parent.parent / "assets"
    header_img = assets_dir / "header_tecpred.png"
    footer_img = assets_dir / "footer_tecpred.png"
    
    # CABEÇALHO
    header = section.header
    if header_img.exists():
        # Remove parágrafos vazios
        for para in header.paragraphs:
            p = para._element
            p.getparent().remove(p)
        
        # Adiciona imagem
        para = header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Alinha à esquerda!
        
        # Remove TODOS os espaçamentos
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        
        # Margem negativa maior para compensar e centralizar a imagem
        para.paragraph_format.left_indent = Cm(-3.0)
        para.paragraph_format.right_indent = Cm(0)
        
        run = para.add_run()
        # Largura total da página + 1cm para cobrir completamente
        run.add_picture(str(header_img), width=Cm(22.0))
    
    # RODAPÉ
    footer = section.footer
    if footer_img.exists():
        # Remove parágrafos vazios
        for para in footer.paragraphs:
            p = para._element
            p.getparent().remove(p)
        
        # Adiciona imagem
        para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Alinha à esquerda!
        
        # Remove TODOS os espaçamentos
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        
        # Margem negativa maior para compensar e centralizar a imagem
        para.paragraph_format.left_indent = Cm(-3.0)
        para.paragraph_format.right_indent = Cm(0)
        
        run = para.add_run()
        # Largura total da página + 1cm para cobrir completamente
        run.add_picture(str(footer_img), width=Cm(22.0))


def add_cover_page(doc: Document, logo_path: str, project_data: dict, memorial_type: str = "telecom"):
    """Adiciona capa ao documento seguindo o modelo padrão.
    
    Estrutura do modelo:
    - Número do memorial (direita, topo)
    - Construtora (centro, negrito)
    - Título do memorial (centro, negrito)
    - Nome do empreendimento (centro, negrito)
    - Endereço (centro)
    - Data (centro, no final)
    
    Args:
        doc: Documento python-docx
        logo_path: Caminho para logo (opcional)
        project_data: Dados do projeto
        memorial_type: Tipo de memorial ("telecom" ou "eletrico")
    """
    from pathlib import Path
    from datetime import datetime
    
    # Número do memorial (direita, topo)
    memo_num = doc.add_paragraph()
    memo_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = memo_num.add_run("MEMORIAL n° XX/2025")  # Ajustar conforme necessário
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    
    # Espaços
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Construtora (centro, negrito)
    construtora = project_data.get("construtora", "")
    if construtora:
        const_para = doc.add_paragraph()
        const_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = const_para.add_run(construtora.upper())
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
    
    # Espaços
    for _ in range(6):
        doc.add_paragraph()
    
    # Título do memorial (centro, negrito) - CONDICIONAL
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if memorial_type == "eletrico":
        empreendimento = project_data.get("empreendimento", "")
        title_text = f"MEMORIAL DESCRITIVO E ESPECIFICAÇÕES TÉCNICAS DAS INSTALAÇÕES ELÉTRICAS DO EDIFÍCIO {empreendimento.upper()}"
    else:
        title_text = "MEMORIAL DESCRITIVO E ESPECIFICAÇÕES TÉCNICAS DAS INSTALAÇÕES DE TELECOMUNICAÇÕES"
    run = title.add_run(title_text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    
    # Espaços
    for _ in range(10):
        doc.add_paragraph()
    
    # Nome do empreendimento (centro, negrito)
    empreendimento = project_data.get("empreendimento", "")
    if empreendimento:
        emp_para = doc.add_paragraph()
        emp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = emp_para.add_run(empreendimento.upper())
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        
        # Endereço (centro, sem negrito)
        endereco = project_data.get("endereco", "")
        if endereco:
            end_para = doc.add_paragraph()
            end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = end_para.add_run(endereco)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
    
    # Espaços até o final
    for _ in range(2):
        doc.add_paragraph()
    
    # Data (centro, final da página)
    data_para = doc.add_paragraph()
    data_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Usa data atual ou do projeto
    carimbo = project_data.get("carimbo", {})
    data = carimbo.get("data", datetime.now().strftime("%d de %B de %Y"))
    run = data_para.add_run(data.upper())
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    
    # Quebra de página
    doc.add_page_break()


def add_table_of_contents(doc: Document):
    """Adiciona sumário conforme modelo.
    
    Estrutura:
    - Título "SUMÁRIO" centralizado
    - Lista de seções com pontos de preenchimento e números de página
    """
    # Título SUMÁRIO (centralizado)
    sumario_title = doc.add_paragraph()
    sumario_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sumario_title.add_run("SUMÁRIO")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    
    doc.add_paragraph()  # Espaço
    
    # Itens do sumário (manualmente, já que python-docx não suporta TOC automático)
    toc_items = [
        ("1. INTRODUÇÃO", False),
        ("2. DADOS DA OBRA", False),
        ("3. NORMAS TÉCNICAS", False),
        ("4. SERVIÇOS CONTEMPLADOS", False),
        ("4.1. SERVIÇO DE VOZ", True),
        ("4.2. SERVIÇO DE DADOS", True),
        ("4.3. SERVIÇO DE VÍDEO", True),
        ("4.4. SERVIÇO DE INTERCOMUNICAÇÃO", True),
        ("4.5. SERVIÇO DE MONITORAMENTO", True),
        ("5. SALA DE MONITORAMENTO (ER / EF)", False),
        ("6. ELEMENTOS PASSIVOS E ATIVOS DA REDE", False),
        ("7. TESTES E ACEITAÇÃO", False),
    ]
    
    for item_text, is_subsection in toc_items:
        para = doc.add_paragraph()
        
        # Indentação para subseções (4.1, 4.2, etc.)
        if is_subsection:
            para.paragraph_format.left_indent = Cm(1.5)
        else:
            para.paragraph_format.left_indent = Cm(0)
        
        # Texto do item (sem pontos de preenchimento)
        run = para.add_run(item_text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    
    # Quebra de página após sumário
    doc.add_page_break()


def add_section_heading(doc: Document, number: str, title: str, level: int = 1):
    """Adiciona título de seção conforme modelo.
    
    Args:
        doc: Documento
        number: Número da seção (ex: "1", "4.1")
        title: Título da seção
        level: Nível (1=Heading 1, 2=Heading 2)
    """
    style = 'Heading 1' if level == 1 else 'Heading 2'
    heading = doc.add_paragraph(style=style)
    
    # Formatação conforme modelo (sem tabulação excessiva)
    text = f"{number}.         {title}" if "." in number and number[0].isdigit() else f"{number}. {title}"
    run = heading.add_run(text)
    
    return heading


def add_body_text(doc: Document, text: str):
    """Adiciona texto de corpo.
    
    Args:
        doc: Documento
        text: Texto a adicionar
    """
    if not text or not text.strip():
        return
    
    # Split por parágrafos
    paragraphs = text.split("\n\n")
    
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        
        # Verifica se é item de lista (começa com -, *, •, ou número)
        if para_text.startswith(("- ", "* ", "• ")) or (len(para_text) > 2 and para_text[0].isdigit() and para_text[1] in ".)-"):
            # Lista
            para = doc.add_paragraph(para_text, style='List Bullet')
        else:
            # Corpo normal
            para = doc.add_paragraph(para_text, style='Normal')


def format_decimal(value: float) -> str:
    """Formata número com vírgula decimal (PT-BR).
    
    Args:
        value: Valor numérico
        
    Returns:
        String formatada
    """
    return f"{value:.2f}".replace(".", ",")











