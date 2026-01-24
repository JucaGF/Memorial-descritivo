"""Sumário fixo para memorial de telecomunicações.

Como as seções do memorial telecom são sempre fixas, usamos um sumário estático
em vez de TOC field dinâmico para evitar problemas de duplicação.
"""

# Seções padrão do memorial telecom (ordem fixa)
TELECOM_TOC_SECTIONS = [
    "1. INTRODUÇÃO",
    "2. DADOS DA OBRA",
    "3. NORMAS TÉCNICAS",
    "4. SERVIÇOS CONTEMPLADOS",
    "4.1. SERVIÇO DE VOZ",
    "4.2. SERVIÇO DE DADOS",
    "4.3. SERVIÇO DE VÍDEO",
    "4.4. SERVIÇO DE INTERCOMUNICAÇÃO",
    "4.5. SERVIÇO DE MONITORAMENTO",
    "5. SALA DE MONITORAMENTO (ER/EF)",
    "6. ELEMENTOS PASSIVOS E ATIVOS DA REDE",
    "7. TESTES E ACEITAÇÃO",
]


def add_fixed_toc_telecom(doc):
    """Adiciona sumário fixo para memorial telecom.
    
    Args:
        doc: Documento python-docx
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Título SUMÁRIO (centralizado)
    sumario_title = doc.add_paragraph()
    sumario_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sumario_title.add_run("SUMÁRIO")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    
    doc.add_paragraph()  # Espaço
    
    # Lista de seções
    for section in TELECOM_TOC_SECTIONS:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Texto da seção
        run = para.add_run(section)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        
        # Adiciona líder pontilhado e número de página como texto
        # (números de página serão adicionados manualmente ou deixados em branco)
        run_dots = para.add_run(" " + "." * 50)
        run_dots.font.name = 'Arial'
        run_dots.font.size = Pt(11)
        
        # Espaço para número de página
        run_page = para.add_run(" __")
        run_page.font.name = 'Arial'
        run_page.font.size = Pt(11)
    
    # Quebra de página após sumário
    doc.add_page_break()
