"""Indexação de memoriais-modelo para retrieval de estilo."""

import re
from pathlib import Path
from typing import List, Dict, Any
import docx  # python-docx

try:
    from langchain_community.vectorstores import FAISS
    from langchain_openai import OpenAIEmbeddings
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False
    FAISS = None
    OpenAIEmbeddings = None
    RecursiveCharacterTextSplitter = None
    Document = None

from memorial_maker.config import settings
from memorial_maker.utils.logging import get_logger

logger = get_logger("rag.index")


class StyleIndexer:
    """Indexa memoriais-modelo para retrieval de estilo/estrutura."""

    def __init__(self):
        """Inicializa indexador."""
        if not LANGCHAIN_AVAILABLE:
            self.embeddings = None
            self.text_splitter = None
            self.vectorstore = None
            return
        
        self.embeddings = OpenAIEmbeddings(
            model=settings.embed_model,
            openai_api_key=settings.openai_api_key,
        )
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", ". ", " "],
        )
        self.vectorstore = None

    def load_doc_file(self, doc_path: Path) -> str:
        """Carrega texto de arquivo DOC/DOCX."""
        try:
            doc = docx.Document(doc_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            logger.info(f"Carregado: {doc_path.name} ({len(text)} chars)")
            return text
        except Exception as e:
            logger.error(f"Erro ao carregar {doc_path.name}: {e}")
            return ""

    def index_models(self, models_dir: Path):
        """Indexa todos os memoriais-modelo de um diretório.
        
        Args:
            models_dir: Diretório com DOC/DOCX
            
        Returns:
            Vectorstore FAISS indexado ou None
        """
        if not LANGCHAIN_AVAILABLE:
            logger.warning("LangChain não está disponível. Indexação desabilitada.")
            return None
        
        from memorial_maker.utils.io_paths import list_models
        
        model_paths = list_models(models_dir)
        logger.info(f"Indexando {len(model_paths)} memoriais-modelo...")
        
        all_docs = []
        
        for model_path in model_paths:
            text = self.load_doc_file(model_path)
            if not text:
                continue
            
            # Split em chunks
            chunks = self.text_splitter.split_text(text)
            
            # Cria documentos com metadata
            for chunk in chunks:
                # Detecta seção
                secao = self._detect_section(chunk)
                
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": model_path.name,
                        "section": secao,
                        "type": "style_reference",
                    }
                )
                all_docs.append(doc)
        
        logger.info(f"Criados {len(all_docs)} chunks de estilo")
        
        # Verifica se há documentos para indexar
        if not all_docs:
            logger.warning("Nenhum documento para indexar, vectorstore não será criado")
            self.vectorstore = None
            return None
        
        # Indexa
        self.vectorstore = FAISS.from_documents(all_docs, self.embeddings)
        logger.info("Indexação concluída")
        
        return self.vectorstore

    def _detect_section(self, chunk: str) -> str:
        """Detecta a qual seção o chunk pertence."""
        chunk_lower = chunk.lower()
        
        # Padrões de seções
        patterns = {
            "introducao": r"1\.\s*introdu[çc][ãa]o",
            "dados_obra": r"2\.\s*dados\s+da\s+obra",
            "normas": r"3\.\s*normas\s+t[ée]cnicas",
            "servicos": r"4\.\s*servi[çc]os\s+contemplados",
            "voz": r"4\.1\.\s*(?:servi[çc]o\s+de\s+)?voz",
            "dados": r"4\.2\.\s*(?:servi[çc]o\s+de\s+)?dados",
            "video": r"4\.3\.\s*(?:servi[çc]o\s+de\s+)?v[íi]deo",
            "intercom": r"4\.4\.\s*(?:servi[çc]o\s+de\s+)?intercomunica[çc][ãa]o",
            "monitoramento": r"4\.5\.\s*(?:servi[çc]o\s+de\s+)?monitoramento",
            "sala": r"5\.\s*sala\s+de\s+monitoramento",
            "passivos": r"6\.\s*elementos\s+passivos",
            "testes": r"7\.\s*testes\s+e\s+aceita[çc][ãa]o",
        }
        
        for secao, pattern in patterns.items():
            if re.search(pattern, chunk_lower):
                return secao
        
        return "geral"

    def retrieve_style_examples(
        self,
        section: str,
        top_k: int = 3,
    ) -> List[str]:
        """Recupera exemplos de estilo para uma seção.
        
        Args:
            section: Nome da seção
            top_k: Número de exemplos
            
        Returns:
            Lista de chunks de exemplo
        """
        if not self.vectorstore:
            logger.warning("Vectorstore não inicializado")
            return []
        
        # Query baseada na seção
        queries = {
            "introducao": "introdução escopo objetivo sistema telecomunicações",
            "dados_obra": "dados da obra empreendimento construtora endereço",
            "normas": "normas técnicas NBR EIA TIA ISO",
            "servicos": "serviços contemplados infraestrutura",
            "voz": "serviço de voz PABX telefone VoIP",
            "dados": "serviço de dados rede estruturada CAT-6 RJ-45 Wi-Fi",
            "video": "serviço de vídeo TV coletiva RG-06 divisores",
            "intercom": "intercomunicação interfone porteiro",
            "monitoramento": "monitoramento CFTV câmeras IP",
            "sala": "sala de monitoramento ER EF rack requisitos",
            "passivos": "elementos passivos ativos materiais patch panel",
            "testes": "testes aceitação certificação",
        }
        
        query = queries.get(section, section)
        
        # Busca com filtro de seção
        try:
            docs = self.vectorstore.similarity_search(
                query,
                k=top_k * 2,  # Busca mais para filtrar
            )
            
            # Filtra por seção ou geral
            filtered = []
            for doc in docs:
                meta_section = doc.metadata.get("section", "")
                if meta_section == section or meta_section == "geral":
                    filtered.append(doc.page_content)
                if len(filtered) >= top_k:
                    break
            
            # Se não achou o suficiente, usa qualquer um
            if len(filtered) < top_k:
                filtered = [doc.page_content for doc in docs[:top_k]]
            
            return filtered
            
        except Exception as e:
            logger.error(f"Erro ao recuperar exemplos para {section}: {e}")
            return []


def index_models(models_dir: Path) -> StyleIndexer:
    """Função conveniente para indexar modelos.
    
    Args:
        models_dir: Diretório com memoriais-modelo
        
    Returns:
        Indexador com vectorstore pronto
    """
    indexer = StyleIndexer()
    indexer.index_models(models_dir)
    return indexer




