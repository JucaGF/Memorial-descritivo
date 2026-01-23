#!/usr/bin/env python3
"""Teste simples para verificar implementações do memorial telecom."""

import sys
from pathlib import Path

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from memorial_maker.writer.docx_styles import add_table_of_contents, add_signature_block
from memorial_maker.rag.generate_sections import SectionGenerator
from memorial_maker.rag.index_style import StyleIndexer


def test_toc_field():
    """Test TOC field implementation."""
    print("=" * 60)
    print("TEST: TOC Field Implementation")
    print("=" * 60)
    
    doc = Document()
    add_table_of_contents(doc)
    
    # Save test document
    output_path = Path("/tmp/test_toc.docx")
    doc.save(str(output_path))
    
    print(f"✓ TOC document created: {output_path}")
    print("  → Open in Word/LibreOffice and right-click to 'Update Field'")
    print("  → Verify dot leaders and page numbers appear")
    print()


def test_signature_block():
    """Test signature block implementation."""
    print("=" * 60)
    print("TEST: Signature Block")
    print("=" * 60)
    
    doc = Document()
    
    # Test with data
    master_data = {
        "obra": {
            "carimbo": {
                "responsavel": "João Silva",
                "crea": "CREA SP 123456"
            }
        }
    }
    
    add_signature_block(doc, master_data)
    
    output_path = Path("/tmp/test_signature.docx")
    doc.save(str(output_path))
    
    print(f"✓ Signature document created: {output_path}")
    print("  → Verify signature line, name, and CREA appear")
    print()
    
    # Test with placeholders
    doc2 = Document()
    add_signature_block(doc2, {})
    
    output_path2 = Path("/tmp/test_signature_placeholder.docx")
    doc2.save(str(output_path2))
    
    print(f"✓ Signature with placeholders: {output_path2}")
    print("  → Verify 'RESPONSÁVEL TÉCNICO' and 'CREA XXXXX' appear")
    print()


def test_building_type_detection():
    """Test building type detection."""
    print("=" * 60)
    print("TEST: Building Type Detection")
    print("=" * 60)
    
    prompts_dir = Path(__fileproject).parent / "memorial_maker" / "rag" / "prompts"
    indexer = StyleIndexer(examples_dir=prompts_dir)
    generator = SectionGenerator(indexer, prompts_dir, memorial_type="telecom")
    
    # Test residencial
    master_data_res = {
        "obra": {
            "empreendimento": "Edifício Residencial Torre A"
        }
    }
    
    building_type = generator._get_building_type(master_data_res)
    print(f"✓ Residencial detection: {building_type}")
    assert building_type == "residencial", f"Expected 'residencial', got '{building_type}'"
    
    # Test corporativo
    master_data_corp = {
        "obra": {
            "tipologia": "Edificação Comercial"
        }
    }
    
    building_type = generator._get_building_type(master_data_corp)
    print(f"✓ Corporativo detection: {building_type}")
    assert building_type == "corporativo", f"Expected 'corporativo', got '{building_type}'"
    
    print()


def test_static_templates():
    """Test static template loading."""
    print("=" * 60)
    print("TEST: Static Template Loading")
    print("=" * 60)
    
    prompts_dir = Path(__file__).parent / "memorial_maker" / "rag" / "prompts"
    indexer = StyleIndexer(examples_dir=prompts_dir)
    generator = SectionGenerator(indexer, prompts_dir, memorial_type="telecom")
    
    # Test template loading
    template_voz_res = generator._load_static_template("s4_1_voz", "residencial")
    print(f"✓ Template s4_1_voz_residencial loaded: {len(template_voz_res) if template_voz_res else 0} chars")
    
    template_voz_corp = generator._load_static_template("s4_1_voz", "corporativo")
    print(f"✓ Template s4_1_voz_corporativo loaded: {len(template_voz_corp) if template_voz_corp else 0} chars")
    
    template_intercom_res = generator._load_static_template("s4_4_intercom", "residencial")
    print(f"✓ Template s4_4_intercom_residencial loaded: {len(template_intercom_res) if template_intercom_res else 0} chars")
    
    template_intercom_corp = generator._load_static_template("s4_4_intercom", "corporativo")
    print(f"✓ Template s4_4_intercom_corporativo loaded: {len(template_intercom_corp) if template_intercom_corp else 0} chars")
    
    observation = generator._load_s5_observation("residencial")
    print(f"✓ Observation s5_observacao_residencial loaded: {len(observation) if observation else 0} chars")
    
    observation_corp = generator._load_s5_observation("corporativo")
    print(f"✓ Observation for corporativo (should be empty): {len(observation_corp)} chars")
    assert len(observation_corp) == 0, "Corporativo should not have observation"
    
    print()


def main():
    """Run all tests."""
    print("\n")
    print("#" * 60)
    print("# MEMORIAL TELECOM - VALIDATION TESTS")
    print("#" * 60)
    print()
    
    try:
        test_toc_field()
        test_signature_block()
        test_building_type_detection()
        test_static_templates()
        
        print("=" * 60)
        print("✓ ALL TESTS PASSED")
        print("=" * 60)
        print()
        print("Next steps:")
        print("1. Review generated DOCX files in /tmp/")
        print("2. Update TOC field in Word/LibreOffice")
        print("3. Compare with reference memorial")
        print("4. Adjust templates if needed based on reference")
        print()
        
    except Exception as e:
        print("=" * 60)
        print(f"✗ TEST FAILED: {e}")
        print("=" * 60)
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
