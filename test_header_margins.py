#!/usr/bin/env python3
"""Teste para entender as margens do cabeçalho/rodapé no Word."""

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Cria documento de teste
doc = Document()

# Configura página
section = doc.sections[0]
section.page_width = Cm(21)  # A4
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.0)
section.top_margin = Cm(3.5)
section.bottom_margin = Cm(2.54)
section.header_distance = Cm(0)
section.footer_distance = Cm(0)

# Adiciona imagem no cabeçalho
header = section.header
header_img = Path("assets/header_tecpred_fullwidth.png")

if header_img.exists():
    # Remove parágrafos default
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    
    # Teste 1: Largura 16.5cm (área interna)
    print("Testando largura 16.5cm (área disponível dentro das margens)...")
    para = header.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    run.add_picture(str(header_img), width=Cm(16.5))
    
    # Salva teste
    output = Path("memorial/test_header_16.5cm.docx")
    output.parent.mkdir(exist_ok=True)
    doc.save(str(output))
    print(f"✓ Salvo: {output}")

print("\n" + "="*60)
print("Abra o arquivo e verifique se a imagem preenche a largura")
print("Se ainda houver espaços, o Word adiciona margens internas")
print("que não podem ser controladas via python-docx")
print("="*60)
